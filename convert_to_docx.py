"""
Convert training_report.md to DOCX format with proper formatting.
Uses python-docx to create a Word document from the markdown.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def parse_and_convert_markdown_to_docx(md_file, output_file):
    """Convert markdown to DOCX with proper formatting."""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines unless between sections
        if not line.strip():
            i += 1
            continue
        
        # Main title (# )
        if re.match(r'^# (?!#)', line):
            title_text = line[2:].strip()
            p = doc.add_heading(title_text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.bold = True
            
        # Section headers (## )
        elif re.match(r'^## (?!#)', line):
            header_text = line[3:].strip()
            p = doc.add_heading(header_text, level=2)
            run = p.runs[0]
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 51, 102)
            
        # Subsection headers (### )
        elif re.match(r'^### (?!#)', line):
            header_text = line[4:].strip()
            p = doc.add_heading(header_text, level=3)
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(51, 51, 51)
            
        # Sub-subsection headers (#### )
        elif re.match(r'^#### ', line):
            header_text = line[5:].strip()
            p = doc.add_heading(header_text, level=4)
            run = p.runs[0]
            run.font.size = Pt(12)
            
        # Horizontal rule (---)
        elif line.strip() == '---':
            doc.add_paragraph()  # Add space
            
        # Image reference
        elif line.strip().startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                doc.add_paragraph(f"[Figure: {alt_text}]", style='Intense Quote')
                # Try to add actual image if it exists
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        pass
        
        # Code block (```)
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(51, 51, 51)
            # Add light gray background effect
            
        # Table (starts with |)
        elif line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Back up one line
            
            if len(table_lines) >= 2:
                # Parse table
                rows = []
                for tline in table_lines:
                    if not re.match(r'^\|[\s\-:]+\|', tline):  # Skip separator line
                        cells = [c.strip() for c in tline.split('|')[1:-1]]
                        rows.append(cells)
                
                if rows:
                    # Create table
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    # Fill table
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Bold header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
        
        # Bullet list (- or *)
        elif re.match(r'^[\-\*] ', line):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            
        # Numbered list (1. or 2. etc)
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            doc.add_paragraph(text, style='List Number')
            
        # Math equations ($$)
        elif line.strip() == '$$':
            eq_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                eq_lines.append(lines[i])
                i += 1
            eq_text = '\n'.join(eq_lines)
            p = doc.add_paragraph(eq_text, style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.italic = True
            run.font.size = Pt(11)
        
        # Regular paragraph
        else:
            if line.strip():
                # Handle inline formatting
                text = line.strip()
                
                # Handle bold (**text**)
                parts = re.split(r'(\*\*.*?\*\*)', text)
                p = doc.add_paragraph()
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        # Handle inline code (`code`)
                        code_parts = re.split(r'(`.*?`)', part)
                        for cpart in code_parts:
                            if cpart.startswith('`') and cpart.endswith('`'):
                                run = p.add_run(cpart[1:-1])
                                run.font.name = 'Consolas'
                                run.font.size = Pt(10)
                            else:
                                p.add_run(cpart)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✓ Converted {md_file} → {output_file}")

if __name__ == "__main__":
    parse_and_convert_markdown_to_docx(
        "training_report.md",
        "training_report.docx"
    )
    print("\nDOCX file created successfully!")
    print("Note: Mathematical equations are shown as text. For proper LaTeX rendering,")
    print("      consider using Pandoc with --mathjax or MathType.")
